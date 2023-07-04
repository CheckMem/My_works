
from docx import Document

document = Document()

document.add_heading('Не забыть взять с собой', 0)
document.add_paragraph('Паспорт')

p = document.add_paragraph('Также надо')
p.add_run('не забыть, ').bold = True

document.add_paragraph('Это будет буллет',style='List Bullet')
document.add_paragraph('Это будет 1',style='List Number')
document.add_paragraph('Это будет 2',style='List Number')
document.add_paragraph('С большой силой приходит большая ответственность',style='Intense Quote')

table = document.add_table( rows=1, cols=2)
header_cells = table.rows[0].cells
header_cells[0].text = '№'
header_cells[1].text = 'Количество'


document.save('text2.docx')





















# from pymorphy2 import MorphAnalyzer
#
# form = MorphAnalyzer().parse('бутылка')[0]
#
# for bottle in reversed(range(9)):
#     print(f'В холодильнике {bottle + 1} {form.make_agree_with_number(bottle + 1).word} пива.')
#     print('Возьмём одну и выпьем.')
#
#     if bottle % 10 == 1 and bottle !=11:
#         remain = 'Осталась'
#     else:
#         remain = 'Осталось'
#     print(f'{remain} {bottle} {form.make_agree_with_number(bottle + 1).word} пива.')











# from pymorphy2 import MorphAnalyzer
#
# form = MorphAnalyzer().parse('бутылка')[0]
#
# for bottle in reversed(range(9)):
#     print(f'В холодильнике {bottle + 1} {form.make_agree_with_number(bottle + 1).word} пива.')
#